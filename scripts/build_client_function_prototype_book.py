from __future__ import annotations

from collections import Counter, defaultdict
from datetime import date
from pathlib import Path
from re import split as re_split

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
EXCEL_PATH = ROOT / "旅游接待管理系统-功能建设清单-完整保留版.xlsx"
OUTPUT_PATH = ROOT / "文档" / "旅游接待管理系统-甲方功能原型图册.docx"

BASE_FONT = "Calibri"
CN_FONT = "Microsoft YaHei"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120

INK = RGBColor(11, 37, 69)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(86, 99, 116)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GREEN = RGBColor(22, 101, 52)
RED = RGBColor(153, 27, 27)
PURPLE = RGBColor(88, 28, 135)

FILL_HEADER = "F2F4F7"
FILL_BLUE = "E8F1FB"
FILL_LIGHT = "F8FAFC"
FILL_PANEL = "F4F6F9"
FILL_GREEN = "ECFDF5"
FILL_WARN = "FFF7ED"
FILL_RED = "FEE2E2"
FILL_PURPLE = "F3E8FF"
FILL_WHITE = "FFFFFF"

MODULE_ORDER = ["客户管理", "采购管理", "销售管理", "计调操作", "财务管理", "数据统计", "企业资料", "系统设置"]

PHASE_FILL = {"P0": FILL_RED, "P1": FILL_BLUE, "P2": FILL_PURPLE}
PHASE_TEXT = {"P0": "P0 一期主线", "P1": "P1 增强/候选", "P2": "P2 后续规划"}

CASE = {
    "customer": "杭州远行国旅",
    "subject": "杭州远行国际旅行社有限公司",
    "team": "HZ20260518-003 西湖宋城二日游",
    "guide": "导游：陈晨",
    "vehicle": "浙A·D2398 37座",
    "supplier": "杭州湖滨酒店 / 西湖车队",
    "order": "订单 S04-001",
    "amount": "¥86,400",
    "payable": "¥52,800",
    "profit": "预计毛利 ¥28,600",
}


def clean(value: object) -> str:
    text = "" if value is None else str(value).strip()
    replacements = {
        "帐款": "账款",
        "应收帐款": "应收账款",
        "应付帐款": "应付账款",
        "帐号": "账号",
        "帐户": "账户",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def split_features(text: str, max_items: int = 8) -> list[str]:
    parts = [clean(x) for x in re_split(r"[、,，;；/]+", clean(text)) if clean(x)]
    seen: list[str] = []
    for item in parts:
        if item not in seen:
            seen.append(item)
    return seen[:max_items] or ["业务信息", "状态", "负责人", "附件"]


def read_excel_items() -> list[dict[str, str]]:
    wb = load_workbook(EXCEL_PATH, read_only=True, data_only=True)
    ws = wb["功能建设清单"]
    headers = [clean(cell.value) for cell in ws[4]]
    items: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=5, values_only=True):
        if not row[0]:
            continue
        data = {headers[i]: clean(row[i]) for i in range(min(len(headers), len(row)))}
        items.append(data)
    return items


def set_run_font(run, *, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_styles(doc: Document):
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = BASE_FONT
        style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = DARK_BLUE
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def setup_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("旅游接待管理系统 | 甲方功能原型图册")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    add_field(footer, "PAGE")
    r = footer.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def add_field(paragraph, field: str):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_para(doc: Document, text: str = "", *, style: str | None = None, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(clean(text))
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, size: float = 8.5, bold: bool = False, color: RGBColor | None = None, fill: str | None = None, align=None):
    if fill:
        shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    r = p.add_run(clean(text))
    set_run_font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_cell_line(cell, text: str, *, size: float = 8.0, bold: bool = False, color: RGBColor | None = None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(clean(text))
    set_run_font(r, size=size, color=color, bold=bold)


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    set_table_margins(table)


def set_table_margins(table, top=80, bottom=80, start=120, end=120):
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = cell_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_table_together(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths: list[int] | None = None):
    widths = widths or [1700, CONTENT_DXA - 1700]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, widths)
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.rows[idx].cells[0], label, bold=True, color=DARK_BLUE, fill=FILL_HEADER, size=8.8)
        set_cell_text(table.rows[idx].cells[1], value, size=8.8)
    return table


def add_callout(doc: Document, title: str, body: str, fill: str = FILL_PANEL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(clean(title))
    set_run_font(r, size=10.2, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    r2 = p2.add_run(clean(body))
    set_run_font(r2, size=9.2, color=INK)
    return table


def classify_item(item: dict[str, str]) -> str:
    text = " ".join(item.values())
    module = item["模块"]
    title = item["优化建设内容"] or item["老系统菜单/功能"]
    menu = item["老系统菜单/功能"]

    if module == "数据统计":
        return "统计类"
    if module == "系统设置":
        return "系统设置类"
    if module == "企业资料":
        if any(k in title + menu for k in ["角色", "部门", "员工"]):
            return "系统设置类"
        return "档案类"
    if module == "客户管理":
        if "授信" in title or "应收" in title or "应收" in menu:
            return "财务类"
        if "合同" in title or "正式主体" in title or "授权" in title:
            return "订单团队类"
        return "档案类"
    if module == "采购管理":
        if any(k in title + menu for k in ["车辆", "排班"]):
            return "计调类"
        return "档案类"
    if module == "销售管理":
        if any(k in title + menu for k in ["费用变更", "成本分摊"]):
            return "财务类"
        if any(k in title + menu for k in ["智能报价", "票务", "门票"]):
            return "计调类"
        return "订单团队类"
    if module == "计调操作":
        if any(k in title + menu for k in ["报账"]):
            return "财务类"
        return "计调类"
    if module == "财务管理":
        return "财务类"
    if any(k in text for k in ["统计", "报表", "进度看板"]):
        return "统计类"
    if any(k in text for k in ["参数", "权限", "日志", "通知", "预警配置", "审批流"]):
        return "系统设置类"
    if any(k in text for k in ["应收", "应付", "收款", "付款", "发票", "备用金", "结算", "成本", "返佣"]):
        return "财务类"
    if item["模块"] == "计调操作" or any(k in text for k in ["计调", "导游", "车辆", "房态", "住宿", "车调", "派车", "报账"]):
        return "计调类"
    if any(k in text for k in ["订单", "团队", "团期", "产品", "拼团", "游客", "报价", "合同"]):
        return "订单团队类"
    return "档案类"


def sample_records(item: dict[str, str], template: str) -> list[list[str]]:
    title = item["优化建设内容"] or item["老系统菜单/功能"]
    if template == "财务类":
        return [
            [CASE["team"], CASE["customer"], CASE["amount"], "待审核", "财务 周敏"],
            ["费用变更 +¥4,800", CASE["customer"], "授信占用", "审批中", "销售 张伟"],
            [CASE["supplier"], "资源应付", CASE["payable"], "待付款", "财务 周敏"],
        ]
    if template == "计调类":
        return [
            ["导游安排", CASE["guide"], "已确认", "无冲突", "计调 刘洋"],
            ["车辆安排", CASE["vehicle"], "待比价", "车价偏高", "车调 王磊"],
            ["住宿安排", "杭州湖滨酒店 18间", "待确认", "房态紧张", "计调 刘洋"],
        ]
    if template == "统计类":
        return [
            ["本月散拼", "286人", "同比 +18%", "已生成", "管理层"],
            [CASE["team"], CASE["profit"], "毛利率 33%", "待复核", "老板"],
            ["杭州远行国旅", "应收 ¥86,400", "账龄 9天", "预警", "财务"],
        ]
    if template == "系统设置类":
        return [
            ["超授信审批", "额度超过 100%", "财务经理→老板", "启用", "管理员"],
            ["成本超预算", "超过 5%", "计调主管→财务", "启用", "管理员"],
            ["订单确认提醒", "确认件缺失", "销售负责人", "启用", "管理员"],
        ]
    if template == "订单团队类":
        return [
            [CASE["order"], CASE["customer"], CASE["team"], CASE["amount"], "待确认"],
            ["拼团 P20260518", "上海春秋门店", CASE["team"], "32人", "已拼团"],
            [title, CASE["customer"], "费用/游客/合同", "待补齐", "审批中"],
        ]
    return [
        [CASE["customer"], "A类客户", "授信 80万", "正常", "销售 张伟"],
        [CASE["supplier"], "协议供应商", "结算月结", "有效", "采购 李娜"],
        [title, "基础资料", "附件已上传", "待确认", "管理员"],
    ]


def business_explanation(item: dict[str, str], template: str) -> str:
    title = item["优化建设内容"] or item["老系统菜单/功能"]
    if template == "档案类":
        return f"{title}是后续下单、排团、结算和统计的基础资料，先把主体、状态、负责人和附件维护清楚，后面业务才能少重复录入。"
    if template == "订单团队类":
        return f"{title}用于承接客户需求和团队执行信息，订单确认后会影响人数、应收、授信、计调安排和利润统计。"
    if template == "计调类":
        return f"{title}让计调能围绕一个团队集中安排资源，及时发现导游、车辆、房态、门票、成本和凭证异常。"
    if template == "财务类":
        return f"{title}把订单收入、资源成本、凭证、收付款和结算状态关联起来，财务能提前看风险，不等团队结束后再补账。"
    if template == "统计类":
        return f"{title}从订单、团队、计调和财务数据自动汇总，让老板按统一口径查看收客、利润、账款和资源使用情况。"
    return f"{title}把关键规则配置化，减少靠人工记忆执行，保证审批、预警、权限和日志可追踪。"


def add_badge_row(cell, item: dict[str, str], template: str):
    cell.text = ""
    labels = [
        f"编号 {item['编号']}",
        item["模块"],
        PHASE_TEXT.get(item["优先级"], item["优先级"]),
        item["建设方式"],
        template,
    ]
    for idx, label in enumerate(labels):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(clean(label))
        color = RED if item["优先级"] == "P0" else PURPLE if item["优先级"] == "P2" else DARK_BLUE
        set_run_font(r, size=8.0, color=color, bold=True)


def add_cover(doc: Document, items: list[dict[str, str]]):
    add_para(doc, "甲方功能原型图册", color=BLUE, bold=True, size=10, after=4)
    add_para(doc, "旅游接待管理系统", style="Title")
    add_para(doc, "基于功能建设清单的业务线框原型", style="Subtitle")
    add_kv_table(
        doc,
        [
            ("文档用途", "用于向甲方说明系统功能范围、页面形态和业务流转，不作为最终视觉 UI。"),
            ("生成依据", "旅游接待管理系统-功能建设清单-完整保留版.xlsx"),
            ("覆盖范围", f"共 {len(items)} 项功能，按 Excel 模块顺序全量覆盖。"),
            ("原型形式", "Word 原生业务线框图，一项一卡，便于现场讲解和后续修改。"),
            ("生成日期", date.today().isoformat()),
        ],
    )
    add_callout(
        doc,
        "阅读方式",
        "先看总流程和岗位视角，再按模块查看功能卡片。每张卡片说明业务背景、页面结构、核心字段，以及该功能在一个团从接单到结算中的作用。",
        fill=FILL_GREEN,
    )


def add_overview(doc: Document, items: list[dict[str, str]]):
    doc.add_heading("一、总览：一个团从接单到结算", level=1)
    steps = [
        ("1 客户与合同", "客户建档、正式主体、销售合同、授信额度", "确认客户能不能下单"),
        ("2 销售接单", "产品/团期、订单、拼团、费用变更、游客名单", "订单确认后生成团队和应收"),
        ("3 计调排团", "导游、车辆、酒店、景区、餐饮、外委地接", "资源安排形成预算成本"),
        ("4 财务审核", "团队审核、应收应付、备用金、导游结算、发票", "收入成本凭证进入账款闭环"),
        ("5 经营统计", "收客、利润、账款、资源采购、导游和收付款汇总", "老板看经营结果和风险"),
    ]
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1872, 1872, 1872, 1872, 1872])
    for idx, (title, body, result) in enumerate(steps):
        cell = table.rows[0].cells[idx]
        shade(cell, FILL_BLUE if idx % 2 == 0 else FILL_GREEN)
        set_cell_text(cell, title, bold=True, color=DARK_BLUE, size=9.0)
        add_cell_line(cell, body, size=7.6, color=DARK)
        add_cell_line(cell, f"结果：{result}", size=7.6, bold=True, color=GREEN)

    doc.add_heading("二、岗位视角：谁看什么", level=1)
    rows = [
        ("老板/管理层", "看订单进度、团队利润、欠款风险、异常团队、经营报表", "业务工作台、利润统计、账款统计、收客统计"),
        ("销售", "维护客户、确认订单、上传确认件、处理费用变更、跟进游客名单", "客户授信、订单管理、费用变更、游客信息"),
        ("计调", "安排导游、车、房、票、餐、外委地接，处理冲突和成本异常", "团队安排、房态库存、车调询价、导游报账"),
        ("财务", "审核团队收入成本、应收应付、备用金、发票、收付款和导游结算", "财务审核、实时应收、应付管理、备用金闭环"),
        ("管理员", "维护员工角色、业务参数、审批预警、消息通知和操作日志", "角色权限、员工管理、业务参数、审批预警、日志"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    set_table_geometry(table, [1500, 4300, CONTENT_DXA - 5800])
    for idx, header in enumerate(["岗位", "主要关心", "相关功能页面"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, fill=FILL_HEADER, size=8.8)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(table.rows[row_idx].cells[col_idx], value, bold=(col_idx == 0), size=8.2, color=DARK_BLUE if col_idx == 0 else DARK)

    counts = Counter(item["优先级"] for item in items)
    module_counts = defaultdict(Counter)
    for item in items:
        module_counts[item["模块"]][item["优先级"]] += 1
    doc.add_heading("三、范围统计", level=1)
    add_callout(doc, "功能覆盖", f"本图册覆盖 Excel 清单全部 {len(items)} 项，其中 P0 一期主线 {counts['P0']} 项，P1 增强/候选 {counts['P1']} 项，P2 后续规划 {counts['P2']} 项。", fill=FILL_PANEL)
    table = doc.add_table(rows=1 + len(MODULE_ORDER), cols=5)
    set_table_geometry(table, [1800, 1200, 1200, 1200, CONTENT_DXA - 5400])
    for idx, header in enumerate(["模块", "P0", "P1", "P2", "甲方阅读重点"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, fill=FILL_HEADER, size=8.8)
    focus = {
        "客户管理": "主体、合同、授信和下单资格",
        "采购管理": "资源、供应商、车辆和外委地接基础",
        "销售管理": "产品、团队、订单、拼团、游客和智能化",
        "计调操作": "团队资源安排、审核、房态、车调、报账",
        "财务管理": "应收应付、备用金、结算、发票和收付款",
        "数据统计": "收客、利润、账款、资源和收付款汇总",
        "企业资料": "部门、员工、角色、导游、合同模板",
        "系统设置": "参数、审批预警、消息、日志",
    }
    for row_idx, module in enumerate(MODULE_ORDER, start=1):
        data = module_counts[module]
        values = [module, str(data["P0"]), str(data["P1"]), str(data["P2"]), focus[module]]
        for col_idx, value in enumerate(values):
            set_cell_text(table.rows[row_idx].cells[col_idx], value, bold=(col_idx == 0), size=8.2)


def add_wireframe_card(doc: Document, item: dict[str, str], index: int):
    template = classify_item(item)
    features = split_features(item["具体功能"])
    title = item["优化建设内容"] or item["老系统菜单/功能"]
    doc.add_heading(f"{index:02d}. {item['模块']} / {title}", level=2)

    meta = doc.add_table(rows=1, cols=4)
    set_table_geometry(meta, [1350, 3300, 2350, CONTENT_DXA - 7000])
    set_cell_text(meta.rows[0].cells[0], f"{item['编号']}", bold=True, color=DARK_BLUE, fill=PHASE_FILL.get(item["优先级"], FILL_HEADER), size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(meta.rows[0].cells[1], f"来源菜单：{item['老系统菜单/功能']}", size=8.3, fill=FILL_LIGHT)
    set_cell_text(meta.rows[0].cells[2], f"{PHASE_TEXT.get(item['优先级'], item['优先级'])} / {item['建设方式']}", bold=True, color=DARK_BLUE, size=8.3, fill=FILL_LIGHT)
    set_cell_text(meta.rows[0].cells[3], template, bold=True, color=GREEN if template in ["档案类", "计调类"] else DARK_BLUE, size=8.3, fill=FILL_LIGHT)

    purpose_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(purpose_table, [1700, CONTENT_DXA - 1700])
    set_cell_text(purpose_table.rows[0].cells[0], "页面作用与功能", bold=True, color=DARK_BLUE, fill=FILL_BLUE, size=8.3)
    set_cell_text(purpose_table.rows[0].cells[1], page_usage_summary(item, template), size=8.1, fill=FILL_BLUE)

    prototype_note = doc.add_table(rows=1, cols=1)
    set_table_geometry(prototype_note, [CONTENT_DXA])
    set_cell_text(
        prototype_note.rows[0].cells[0],
        "页面原型示意：下方展示系统页面的主操作区，用于确认左侧怎么查询、中间看什么列表、右侧展示哪些详情，以及状态如何联动。",
        bold=True,
        color=DARK_BLUE,
        fill=FILL_PANEL,
        size=8.2,
    )

    wire = doc.add_table(rows=4, cols=4)
    set_table_geometry(wire, [2150, 2450, 2450, CONTENT_DXA - 7050])
    keep_table_together(wire)
    for col in range(4):
        set_cell_text(wire.rows[0].cells[col], wire_header(template, col), bold=True, color=DARK_BLUE, fill=FILL_BLUE, size=7.8, align=WD_ALIGN_PARAGRAPH.CENTER)

    fill_filters(wire.rows[1].cells[0], item, features, template)
    fill_main_list(wire.rows[1].cells[1], item, features, template)
    fill_detail_panel(wire.rows[1].cells[2], item, features, template)
    fill_status_panel(wire.rows[1].cells[3], item, template)

    set_cell_text(wire.rows[2].cells[0], "核心字段/操作", bold=True, color=DARK_BLUE, fill=FILL_HEADER, size=7.8)
    field_cell = wire.rows[2].cells[1].merge(wire.rows[2].cells[3])
    field_cell.text = ""
    for idx, feature in enumerate(features[:8], start=1):
        add_cell_line(field_cell, f"{idx}. {feature}", size=7.4)

    set_cell_text(wire.rows[3].cells[0], "业务说明", bold=True, color=DARK_BLUE, fill=FILL_GREEN, size=7.8)
    explanation_cell = wire.rows[3].cells[1].merge(wire.rows[3].cells[3])
    set_cell_text(explanation_cell, business_explanation(item, template), size=7.8, fill=FILL_GREEN)

    add_para(doc, after=3)


def page_usage_summary(item: dict[str, str], template: str) -> str:
    title = item["优化建设内容"] or item["老系统菜单/功能"]
    features = "、".join(split_features(item["具体功能"], 4))
    delivery = clean(item["交付说明"]).rstrip("。")
    if template == "档案类":
        return f"页面用于维护{features}等基础信息，并支撑后续下单、排团、结算和统计。{delivery}。"
    if template == "订单团队类":
        return f"页面用于管理{features}等业务信息，承接销售接单、合同确认和团队执行流转。{delivery}。"
    if template == "计调类":
        return f"页面用于安排{features}等计调资源，集中查看确认状态、资源冲突和成本变化。{delivery}。"
    if template == "财务类":
        return f"页面用于管理{features}等财务事项，联动订单、团队、凭证、收付款和账款状态。{delivery}。"
    if template == "统计类":
        return f"页面用于汇总{features}等经营数据，支持按客户、团队、资源和财务口径查看结果。{delivery}。"
    return f"页面用于配置{features}等系统规则，支撑权限、审批、预警、消息和日志留痕。{delivery}。"


def wire_header(template: str, col: int) -> str:
    headers = {
        "档案类": ["左侧查询区", "中间档案列表", "右侧详情区", "状态联动区"],
        "订单团队类": ["左侧查询区", "中间订单/团队列表", "右侧订单明细", "流程状态区"],
        "计调类": ["左侧团队查询", "中间资源安排", "右侧异常/成本", "审核状态区"],
        "财务类": ["左侧账款查询", "中间账款/审核列表", "右侧金额明细", "凭证状态区"],
        "统计类": ["左侧统计条件", "中间指标看板", "右侧图表区域", "明细钻取区"],
        "系统设置类": ["左侧配置查询", "中间规则列表", "右侧规则详情", "权限日志区"],
    }
    return headers.get(template, headers["档案类"])[col]


def fill_filters(cell, item: dict[str, str], features: list[str], template: str):
    set_cell_text(cell, "页面左侧：查询条件", bold=True, color=DARK_BLUE, fill=FILL_LIGHT, size=7.6)
    common = {
        "档案类": ["名称/编号", "分类/状态", "负责人", "更新时间"],
        "订单团队类": ["客户", "团队日期", "订单状态", "销售负责人"],
        "计调类": ["团队号", "出团日期", "资源状态", "计调负责人"],
        "财务类": ["客户/供应商", "团队号", "账款状态", "账龄/日期"],
        "统计类": ["月份", "渠道/客户", "团队类型", "业务员"],
        "系统设置类": ["规则名称", "启停状态", "适用模块", "修改人"],
    }.get(template, ["关键词", "状态", "负责人"])
    for field in common:
        add_cell_line(cell, f"筛选项：{field}", size=7.2, color=MUTED)
    add_cell_line(cell, "页面按钮：查询 / 重置 / 导出", size=7.2, bold=True, color=DARK_BLUE)


def fill_main_list(cell, item: dict[str, str], features: list[str], template: str):
    set_cell_text(cell, "页面中间：列表展示", bold=True, color=DARK_BLUE, fill=FILL_LIGHT, size=7.6)
    rows = sample_records(item, template)
    for row in rows:
        add_cell_line(cell, " | ".join(row[:3]), size=7.0)
    add_cell_line(cell, action_text(template), size=7.2, bold=True, color=DARK_BLUE)


def fill_detail_panel(cell, item: dict[str, str], features: list[str], template: str):
    set_cell_text(cell, "页面右侧：详情内容", bold=True, color=DARK_BLUE, fill=FILL_LIGHT, size=7.6)
    if template == "财务类":
        detail = [f"客户：{CASE['customer']}", f"应收：{CASE['amount']}", f"应付：{CASE['payable']}", "凭证：已上传 2 份"]
    elif template == "计调类":
        detail = [f"团队：{CASE['team']}", f"导游：{CASE['guide']}", f"车辆：{CASE['vehicle']}", "住宿：18间待确认"]
    elif template == "统计类":
        detail = ["收客：286人", "毛利：54.8万", "欠款客户：4户", "异常团队：3个"]
    elif template == "系统设置类":
        detail = ["适用模块：销售/财务", "触发条件：超阈值", "审批节点：主管/财务/老板", "日志：保留"]
    elif template == "订单团队类":
        detail = [f"客户：{CASE['customer']}", f"团队：{CASE['team']}", f"金额：{CASE['amount']}", "确认件：待上传"]
    else:
        detail = [f"名称：{features[0] if features else item['优化建设内容']}", f"关联客户：{CASE['customer']}", "状态：正常", "附件：营业执照/合同"]
    for line in detail:
        add_cell_line(cell, line, size=7.2)


def fill_status_panel(cell, item: dict[str, str], template: str):
    set_cell_text(cell, "页面右侧：状态联动", bold=True, color=DARK_BLUE, fill=FILL_LIGHT, size=7.6)
    status_map = {
        "档案类": ["状态：启用/停用", "联动：订单/合同/账款", "提醒：资料缺失"],
        "订单团队类": ["状态：待确认→已确认", "联动：应收/授信/计调", "提醒：确认件缺失"],
        "计调类": ["状态：待安排→可审核", "联动：成本/应付/报账", "提醒：冲突/超预算"],
        "财务类": ["状态：待审→已通过", "联动：收款/付款/发票", "提醒：账龄/凭证"],
        "统计类": ["状态：已生成/待校验", "联动：订单/财务数据", "提醒：口径异常"],
        "系统设置类": ["状态：启用/停用", "联动：审批/通知/日志", "提醒：规则冲突"],
    }
    for line in status_map.get(template, status_map["档案类"]):
        add_cell_line(cell, line, size=7.2)
    add_cell_line(cell, f"交付：{short_delivery(item['交付说明'])}", size=7.0, color=GREEN, bold=True)


def action_text(template: str) -> str:
    actions = {
        "档案类": "页面按钮：详情 / 编辑 / 导出",
        "订单团队类": "页面按钮：详情 / 编辑 / 确认订单",
        "计调类": "页面按钮：安排资源 / 提交审核",
        "财务类": "页面按钮：审核 / 收付款 / 查看凭证",
        "统计类": "页面按钮：查看明细 / 导出报表",
        "系统设置类": "页面按钮：新增规则 / 编辑 / 启停",
    }
    return actions.get(template, "页面按钮：详情 / 编辑 / 导出")


def short_delivery(text: str) -> str:
    text = clean(text)
    return text if len(text) <= 34 else text[:32] + "..."


def add_module_sections(doc: Document, items: list[dict[str, str]]):
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for item in items:
        grouped[item["模块"]].append(item)

    card_index = 1
    for module in MODULE_ORDER:
        doc.add_page_break()
        module_items = grouped[module]
        doc.add_heading(f"{module}", level=1)
        add_callout(
            doc,
            f"{module}阅读提示",
            module_intro(module, module_items),
            fill=FILL_PANEL,
        )
        for item in module_items:
            add_wireframe_card(doc, item, card_index)
            card_index += 1


def module_intro(module: str, items: list[dict[str, str]]) -> str:
    counts = Counter(item["优先级"] for item in items)
    intro = {
        "客户管理": "客户管理说明“客户是谁、用哪个正式主体签合同开票、能不能继续下单”。",
        "采购管理": "采购管理为计调准备可用资源、供应商、车辆和采购价格，不直接替代每日排团。",
        "销售管理": "销售管理承接客户需求，形成产品、团期、团队、订单、游客、费用变更和后续智能化能力。",
        "计调操作": "计调操作围绕一个团队集中安排资源、审核成本和处理导游报账。",
        "财务管理": "财务管理把订单收入、资源成本、备用金、发票、收付款和结算统一成账款闭环。",
        "数据统计": "数据统计把前面业务数据按统一口径汇总给老板和管理人员看。",
        "企业资料": "企业资料维护组织、人员、角色、导游、合同模板等系统运行基础。",
        "系统设置": "系统设置维护审批、预警、消息、日志和业务参数，让规则可配置可追踪。",
    }[module]
    return f"{intro} 本模块共 {len(items)} 项：P0 {counts['P0']} 项，P1 {counts['P1']} 项，P2 {counts['P2']} 项。"


def validate_doc_source(items: list[dict[str, str]]):
    if len(items) != 67:
        raise ValueError(f"Expected 67 items from Excel, got {len(items)}")
    missing = [item for item in items if not item.get("编号") or not item.get("模块") or not item.get("优化建设内容")]
    if missing:
        raise ValueError(f"Missing required fields: {missing[:3]}")
    actual_modules = list(dict.fromkeys(item["模块"] for item in items))
    if actual_modules != MODULE_ORDER:
        raise ValueError(f"Unexpected module order: {actual_modules}")


def build_doc():
    items = read_excel_items()
    validate_doc_source(items)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_document(doc)
    configure_styles(doc)
    add_cover(doc, items)
    add_overview(doc, items)
    add_module_sections(doc, items)

    props = doc.core_properties
    props.title = "旅游接待管理系统-甲方功能原型图册"
    props.subject = "甲方功能原型图册"
    props.author = "Codex"
    props.comments = "Generated from 旅游接待管理系统-功能建设清单-完整保留版.xlsx only; no frontend code or screenshots reused."
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
