from __future__ import annotations

from collections import Counter
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
WORKBOOK_PATH = ROOT / "旅游接待管理系统-功能建设清单-完整保留版.xlsx"
SCREENSHOT_DIR = ROOT / "output" / "原型确认截图"
DOCX_PATH = ROOT / "文档" / "旅游接待管理系统-原型与方案确认书.docx"

ACCENT = RGBColor(46, 116, 181)
DARK_ACCENT = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 102, 112)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WARNING_FILL = "FFF4DE"
WHITE = RGBColor(255, 255, 255)

BASE_FONT = "Calibri"
CN_FONT = "Microsoft YaHei"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


CORE_SCREENSHOTS = [
    {
        "file": "01-业务工作台.png",
        "title": "业务工作台",
        "route": "/workspace",
        "points": [
            "确认首页是否覆盖待确认订单、待排团、计调异常、应收预警、备用金审批和团队利润。",
            "确认快捷入口是否适合现场岗位顺序：客户授信、订单确认、团队安排、计调审核、财务审核、统计看板。",
            "确认 P0/P1/P2 阶段说明是否能作为商务范围边界。",
        ],
    },
    {
        "file": "02-客户授信.png",
        "title": "客户授信与实时应收",
        "route": "/customer/credit",
        "points": [
            "订单确认和费用变更是否立即占用客户授信额度。",
            "超限时是直接拦截、转审批，还是允许指定角色放行。",
            "收款完成后是否自动释放授信占用，并同步应收账龄。",
        ],
    },
    {
        "file": "03-订单确认.png",
        "title": "订单确认",
        "route": "/sales/order",
        "points": [
            "订单确认是否必须上传客户确认件或合同附件。",
            "散拼订单、整团订单、子订单是否分别生成应收记录。",
            "取消订单、转团、费用变更是否释放或重新占用授信。",
        ],
    },
    {
        "file": "04-计调安排.png",
        "title": "团队安排总控",
        "route": "/dispatch/team-arrange",
        "points": [
            "导游、住宿、用车、景区、餐饮、购物、地接是否都在总控页体现完成状态。",
            "资源安排产生的预算成本是否进入计调审核和财务审核。",
            "车辆冲突、房态不足、成本超预算等异常是否需要预警和审批。",
        ],
    },
    {
        "file": "05-财务审核.png",
        "title": "财务团队审核",
        "route": "/finance/team-audit",
        "points": [
            "财务审核是否前置查看团队应收、应付、预算成本、实际成本和毛利。",
            "计调退回、补凭证、成本超预算的处理责任人是否明确。",
            "审核通过后是否自动生成应收应付和后续收付款任务。",
        ],
    },
    {
        "file": "06-实时应收.png",
        "title": "实时应收管理",
        "route": "/finance/receivable",
        "points": [
            "订单确认即形成应收快照的业务口径是否确认。",
            "费用变更、收款抵扣、授信释放是否需要全程留痕。",
            "账龄、催收、客户授信占用是否进入管理预警。",
        ],
    },
    {
        "file": "07-收客统计.png",
        "title": "收客统计",
        "route": "/statistics/reception",
        "points": [
            "有效人数仅统计成人和儿童，房差及附加项不计入人头。",
            "渠道、客户、团队、业务员的收客归属口径是否确认。",
            "该口径是否作为后续绩效和经营报表基础。",
        ],
    },
]


PROCESS_STEPS = [
    ("客户/采购", "客户主档、正式主体、客户合同、供应商、资源和采购价形成业务基础。"),
    ("销售", "产品、团期、订单、拼团、费用变更、电子合同和游客名单进入销售主流程。"),
    ("计调", "团队安排总控分配导游、住宿、用车、景区、餐饮、地接并形成预算成本。"),
    ("财务", "财务审核、应收应付、收付款、备用金、导游结算和发票形成资金闭环。"),
    ("统计", "团队进度、收客、利润、账款、资源采购等报表使用统一口径输出。"),
]


def sanitize(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return (
        text.replace("帐款", "账款")
        .replace("帐户", "账户")
        .replace("帐号", "账号")
        .replace("应收帐款", "应收账款")
        .replace("应付帐款", "应付账款")
    )


def set_run_font(
    run,
    *,
    name: str = BASE_FONT,
    east_asia: str = CN_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    font = style.font
    font.name = BASE_FONT
    font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    set_style_font(title, size=23, color=INK, bold=True)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, size=12, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, size=16, color=ACCENT, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=13, color=ACCENT, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_ACCENT, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles.add_style("Figure Caption", 1)
    set_style_font(caption, size=9, color=MUTED)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)

    small = doc.styles.add_style("Small Table Text", 1)
    set_style_font(small, size=8.4)
    small.paragraph_format.space_after = Pt(0)
    small.paragraph_format.line_spacing = 1.05


def section_setup(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("旅游接待管理系统 | 商务确认原型")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number_field(footer)
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
    size: float | None = None,
    align: int | None = None,
    before: float | None = None,
    after: float | None = None,
) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        run = p.add_run(sanitize(text))
        set_run_font(run, size=size, color=color, bold=bold)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.2, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(sanitize(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for m, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Iterable[int], *, indent_dxa: int = TABLE_INDENT_DXA) -> None:
    widths = list(widths_dxa)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
    set_cell_margins(table)


def add_metadata_table(doc: Document) -> None:
    rows = [
        ("项目名称", "旅游接待管理系统商务确认原型"),
        ("文档用途", "用于甲方现场讲解、模块范围确认和一期开发边界确认"),
        ("原型地址", "http://127.0.0.1:5667/workspace（本地演示环境）"),
        ("依据文件", "旅游接待管理系统-功能建设清单-完整保留版.xlsx"),
        ("阶段口径", "P0：一期主线；P1：一期增强/二期候选；P2：后续智能化规划"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1600, CONTENT_WIDTH_DXA - 1600])
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], LIGHT_FILL)
        set_cell_text(row.cells[0], label, bold=True, size=9.5, color=DARK_ACCENT)
        set_cell_text(row.cells[1], value, size=9.5)
    add_paragraph(doc, after=6)


def add_callout(doc: Document, title: str, body: str, *, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(sanitize(title))
    set_run_font(run, size=10.5, color=DARK_ACCENT, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(sanitize(body))
    set_run_font(run, size=9.8, color=INK)
    add_paragraph(doc, after=4)


def add_process_table(doc: Document) -> None:
    table = doc.add_table(rows=1 + len(PROCESS_STEPS), cols=3)
    set_table_geometry(table, [1200, 1800, CONTENT_WIDTH_DXA - 3000])
    headers = ["序号", "流程阶段", "原型表达"]
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_FILL)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9.5, color=DARK_ACCENT)
    for i, (stage, desc) in enumerate(PROCESS_STEPS, start=1):
        cells = table.rows[i].cells
        set_cell_text(cells[0], str(i), bold=True, size=9.5)
        set_cell_text(cells[1], stage, bold=True, size=9.5, color=INK)
        set_cell_text(cells[2], desc, size=9.3)
    add_paragraph(doc, after=6)


def read_scope_rows() -> list[dict[str, str]]:
    wb = load_workbook(WORKBOOK_PATH, read_only=True, data_only=True)
    ws = wb["功能建设清单"]
    rows: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=5, values_only=True):
        no, module, old_menu, optimization, functions, problem, build_mode, priority, delivery = row[:9]
        if not no:
            continue
        phase = sanitize(priority)
        rows.append(
            {
                "编号": sanitize(no),
                "模块": sanitize(module),
                "功能": sanitize(optimization or old_menu),
                "来源菜单": sanitize(old_menu),
                "建设方式": sanitize(build_mode),
                "阶段": phase,
                "确认状态": "本次重点确认" if phase == "P0" else "范围确认",
                "备注": sanitize(delivery),
            }
        )
    return rows


def phase_summary(rows: list[dict[str, str]]) -> str:
    counts = Counter(row["阶段"] for row in rows)
    return f"本清单共 {len(rows)} 项：P0 一期主线 {counts.get('P0', 0)} 项，P1 一期增强/二期候选 {counts.get('P1', 0)} 项，P2 后续智能化规划 {counts.get('P2', 0)} 项。"


def add_scope_table(doc: Document, rows: list[dict[str, str]]) -> None:
    add_paragraph(doc, phase_summary(rows), color=INK, bold=True, after=6)
    table = doc.add_table(rows=1 + len(rows), cols=6)
    set_table_geometry(table, [800, 1150, 2100, 950, 1250, CONTENT_WIDTH_DXA - 6250])
    headers = ["编号", "模块", "功能名称", "阶段", "建设方式", "备注/交付说明"]
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_FILL)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=8.6, color=DARK_ACCENT)
    for row_index, item in enumerate(rows, start=1):
        cells = table.rows[row_index].cells
        values = [
            item["编号"],
            item["模块"],
            item["功能"],
            stage_label(item["阶段"]),
            item["建设方式"],
            item["备注"],
        ]
        for col_index, value in enumerate(values):
            set_cell_text(cells[col_index], value, size=7.8 if col_index == 5 else 8.2)
        if item["阶段"] == "P0":
            shade_cell(cells[3], "FDECEC")
        elif item["阶段"] == "P1":
            shade_cell(cells[3], "EAF3FF")
        else:
            shade_cell(cells[3], "F2EAFE")
    add_paragraph(doc, after=6)


def stage_label(phase: str) -> str:
    return {
        "P0": "P0 一期主线",
        "P1": "P1 一期增强/二期候选",
        "P2": "P2 后续智能化规划",
    }.get(phase, phase)


def add_screenshot_section(doc: Document, item: dict[str, object], index: int) -> None:
    doc.add_heading(f"{index}. {item['title']}", level=2)
    add_paragraph(doc, f"原型入口：{item['route']}", color=MUTED, size=9.2, after=4)
    image_path = SCREENSHOT_DIR / str(item["file"])
    if image_path.exists():
        picture = doc.add_picture(str(image_path), width=Inches(6.35))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
        add_paragraph(doc, f"图 {index}：{item['title']} 原型截图", style="Figure Caption", align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_callout(doc, "截图缺失", f"未找到截图文件：{image_path.name}", fill=WARNING_FILL)

    table = doc.add_table(rows=1 + len(item["points"]), cols=2)
    set_table_geometry(table, [1100, CONTENT_WIDTH_DXA - 1100])
    shade_cell(table.rows[0].cells[0], LIGHT_FILL)
    shade_cell(table.rows[0].cells[1], LIGHT_FILL)
    set_cell_text(table.rows[0].cells[0], "序号", bold=True, size=8.8, color=DARK_ACCENT)
    set_cell_text(table.rows[0].cells[1], "甲方确认点", bold=True, size=8.8, color=DARK_ACCENT)
    for row_index, point in enumerate(item["points"], start=1):
        set_cell_text(table.rows[row_index].cells[0], str(row_index), bold=True, size=8.8)
        set_cell_text(table.rows[row_index].cells[1], str(point), size=8.7)
    add_paragraph(doc, after=6)


def add_confirmation_page(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("甲方确认页", level=1)
    add_callout(
        doc,
        "确认说明",
        "本页用于确认原型表达、一期范围、后续规划和需调整意见。签字或批注意见完成后，可作为后续详细设计与开发排期依据。",
    )

    table = doc.add_table(rows=8, cols=2)
    set_table_geometry(table, [2200, CONTENT_WIDTH_DXA - 2200])
    fields = [
        ("确认单位", ""),
        ("确认范围", "Web 可点击原型、P0 一期主线流程、P1/P2 后续范围、模块菜单清单"),
        ("确认结论", "□ 确认通过    □ 修改后通过    □ 需重新评审"),
        ("修改意见", ""),
        ("甲方确认人", ""),
        ("职务/部门", ""),
        ("确认日期", ""),
        ("备注", ""),
    ]
    for row, (label, value) in zip(table.rows, fields):
        shade_cell(row.cells[0], LIGHT_FILL)
        set_cell_text(row.cells[0], label, bold=True, size=9.5, color=DARK_ACCENT)
        set_cell_text(row.cells[1], value, size=9.5)
        if label in {"修改意见", "备注"}:
            for _ in range(3):
                p = row.cells[1].add_paragraph()
                p.paragraph_format.space_after = Pt(8)
    add_paragraph(doc, after=12)
    add_paragraph(doc, "签字/盖章：________________________", size=11, bold=True, after=12)
    add_paragraph(doc, "日期：________ 年 ________ 月 ________ 日", size=11, bold=True)


def build_doc() -> None:
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section_setup(doc)
    configure_styles(doc)

    add_paragraph(doc, "商务确认文件", color=ACCENT, bold=True, size=10, after=4)
    add_paragraph(doc, "旅游接待管理系统", style="Title")
    add_paragraph(doc, "原型与方案确认书", style="Subtitle")
    add_metadata_table(doc)
    add_callout(
        doc,
        "交付目标",
        "本确认书配合 Web 高保真原型使用，用于在商务阶段让甲方确认业务流程、菜单范围、一期开发边界、字段与规则口径。原型为静态样例数据，不连接真实生产系统。",
    )

    doc.add_heading("一、项目背景与原型说明", level=1)
    add_paragraph(
        doc,
        "本次原型基于既有 Vben + Ant Design Vue 后台工程扩展，围绕旅游地接业务从客户授信、订单确认、团队安排、计调审核、财务审核、应收应付到经营统计的闭环进行表达。所有菜单均按功能建设清单进行覆盖，P0 页面优先做成可讲解的高保真业务样例，P1/P2 页面保留真实字段、阶段边界和确认点。",
    )
    add_callout(
        doc,
        "阶段口径",
        "P0 为一期开发主线；P1 为一期增强或二期候选；P2 为后续智能化规划。本文档统一使用“账款”“账号”等术语。",
    )

    doc.add_heading("二、总体业务流程", level=1)
    add_process_table(doc)
    add_callout(
        doc,
        "核心闭环",
        "客户与采购资料提供基础口径，销售订单触发授信和应收，计调安排形成预算成本，财务审核生成账款和结算任务，统计模块输出经营与风控报表。",
    )

    doc.add_heading("三、一期核心流程原型截图与确认点", level=1)
    for index, item in enumerate(CORE_SCREENSHOTS, start=1):
        add_screenshot_section(doc, item, index)

    rows = read_scope_rows()
    doc.add_heading("四、全菜单模块范围表", level=1)
    add_scope_table(doc, rows)

    doc.add_heading("五、演示与验收口径", level=1)
    acceptance = [
        ("菜单覆盖", "Web 原型包含客户、采购、销售、计调、财务、统计、企业资料、系统设置 8 类业务菜单，合计 77 个原型页面。"),
        ("主流程链路", "现场讲解建议按：工作台 → 客户授信 → 产品/团队 → 订单 → 计调安排 → 财务审核 → 应收/应付 → 收客统计。"),
        ("交互范围", "按钮、抽屉、状态、筛选和表格用于演示业务意图；不实现真实保存、审批、支付、导入或接口调用。"),
        ("确认输出", "甲方需确认字段、状态、审批规则、统计口径和一期/后续边界；确认后进入详细设计和开发排期。"),
    ]
    table = doc.add_table(rows=1 + len(acceptance), cols=2)
    set_table_geometry(table, [1800, CONTENT_WIDTH_DXA - 1800])
    shade_cell(table.rows[0].cells[0], LIGHT_FILL)
    shade_cell(table.rows[0].cells[1], LIGHT_FILL)
    set_cell_text(table.rows[0].cells[0], "验收项", bold=True, color=DARK_ACCENT)
    set_cell_text(table.rows[0].cells[1], "说明", bold=True, color=DARK_ACCENT)
    for i, (label, body) in enumerate(acceptance, start=1):
        set_cell_text(table.rows[i].cells[0], label, bold=True, size=9.3)
        set_cell_text(table.rows[i].cells[1], body, size=9.2)

    add_confirmation_page(doc)

    core_props = doc.core_properties
    core_props.title = "旅游接待管理系统-原型与方案确认书"
    core_props.subject = "商务确认原型"
    core_props.author = "Codex"
    core_props.comments = f"Generated on {date.today().isoformat()} from Web prototype screenshots and functional scope workbook."
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
