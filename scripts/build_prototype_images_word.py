from __future__ import annotations

from pathlib import Path
from re import split as re_split

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
EXCEL_PATH = ROOT / "旅游接待管理系统-功能建设清单-完整保留版.xlsx"
IMAGE_DIR = ROOT / "output" / "甲方功能原型页面图"
OUTPUT_PATH = ROOT / "文档" / "旅游接待管理系统-甲方功能原型页面图册.docx"

BASE_FONT = "Calibri"
CN_FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(86, 99, 116)
DARK = RGBColor(17, 24, 39)
GREEN = RGBColor(22, 101, 52)

MODULE_ORDER = ["客户管理", "采购管理", "销售管理", "计调操作", "财务管理", "数据统计", "企业资料", "系统设置"]
PHASE_TEXT = {"P0": "P0 一期主线", "P1": "P1 增强/候选", "P2": "P2 后续规划"}


def clean(value: object) -> str:
    text = "" if value is None else str(value).strip()
    for old, new in {
        "帐款": "账款",
        "应收帐款": "应收账款",
        "应付帐款": "应付账款",
        "帐号": "账号",
        "帐户": "账户",
    }.items():
        text = text.replace(old, new)
    return text


def split_features(text: str, limit: int = 4) -> list[str]:
    parts = [clean(x) for x in re_split(r"[、,，;；/]+", clean(text)) if clean(x)]
    result = []
    for item in parts:
        if item not in result:
            result.append(item)
    return result[:limit] or ["业务信息", "状态", "负责人", "附件"]


def read_items() -> list[dict[str, str]]:
    wb = load_workbook(EXCEL_PATH, read_only=True, data_only=True)
    ws = wb["功能建设清单"]
    headers = [clean(c.value) for c in ws[4]]
    items = []
    for row in ws.iter_rows(min_row=5, values_only=True):
        if not row[0]:
            continue
        items.append({headers[i]: clean(row[i]) for i in range(9)})
    return items


def safe_name(text: str) -> str:
    text = clean(text)
    for ch in '/\\:*?"<>| ':
        text = text.replace(ch, "")
    return text[:28]


def image_path_for(item: dict[str, str], idx: int) -> Path:
    title = item["优化建设内容"] or item["老系统菜单/功能"]
    return IMAGE_DIR / f"{idx:02d}-{safe_name(item['模块'])}-{safe_name(title)}.png"


def classify(item: dict[str, str]) -> str:
    text = " ".join(item.values())
    module = item["模块"]
    title = item["优化建设内容"] or item["老系统菜单/功能"]
    menu = item["老系统菜单/功能"]

    if module == "数据统计":
        return "统计类"
    if module == "系统设置":
        return "系统设置类"
    if module == "企业资料":
        if any(k in title + menu for k in ["角色", "部门", "员工"]):
            return "系统设置类"
        return "档案类"
    if module == "客户管理":
        if "授信" in title or "应收" in title or "应收" in menu:
            return "财务类"
        if "合同" in title or "正式主体" in title or "授权" in title:
            return "订单团队类"
        return "档案类"
    if module == "采购管理":
        if any(k in title + menu for k in ["车辆", "排班"]):
            return "计调类"
        return "档案类"
    if module == "销售管理":
        if any(k in title + menu for k in ["费用变更", "成本分摊"]):
            return "财务类"
        if any(k in title + menu for k in ["智能报价", "票务", "门票"]):
            return "计调类"
        return "订单团队类"
    if module == "计调操作":
        if any(k in title + menu for k in ["报账"]):
            return "财务类"
        return "计调类"
    if module == "财务管理":
        return "财务类"
    if any(k in text for k in ["统计", "报表", "进度看板"]):
        return "统计类"
    if any(k in text for k in ["参数", "权限", "日志", "通知", "预警配置", "审批流"]):
        return "系统设置类"
    if any(k in text for k in ["应收", "应付", "收款", "付款", "发票", "备用金", "结算", "成本", "返佣"]):
        return "财务类"
    if any(k in text for k in ["计调", "导游", "车辆", "房态", "住宿", "车调", "派车", "报账"]):
        return "计调类"
    if any(k in text for k in ["订单", "团队", "团期", "产品", "拼团", "游客", "报价", "合同"]):
        return "订单团队类"
    return "档案类"


def page_usage_summary(item: dict[str, str]) -> str:
    kind = classify(item)
    features = "、".join(split_features(item["具体功能"], 4))
    delivery = clean(item["交付说明"]).rstrip("。")
    if kind == "档案类":
        return f"页面用于维护{features}等基础信息，并支撑后续下单、排团、结算和统计。{delivery}。"
    if kind == "订单团队类":
        return f"页面用于管理{features}等业务信息，承接销售接单、合同确认和团队执行流转。{delivery}。"
    if kind == "计调类":
        return f"页面用于安排{features}等计调资源，集中查看确认状态、资源冲突和成本变化。{delivery}。"
    if kind == "财务类":
        return f"页面用于管理{features}等财务事项，联动订单、团队、凭证、收付款和账款状态。{delivery}。"
    if kind == "统计类":
        return f"页面用于汇总{features}等经营数据，支持按客户、团队、资源和财务口径查看结果。{delivery}。"
    return f"页面用于配置{features}等系统规则，支撑权限、审批、预警、消息和日志留痕。{delivery}。"


def set_run_font(run, *, size=None, color=None, bold=None):
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.28)
    section.bottom_margin = Inches(0.28)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.header_distance = Inches(0.16)
    section.footer_distance = Inches(0.16)

    for style_name in ["Normal", "Title", "Subtitle", "Heading 1"]:
        style = doc.styles[style_name]
        style.font.name = BASE_FONT
        style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Title"].font.size = Pt(24)
    doc.styles["Title"].font.bold = True
    doc.styles["Title"].font.color.rgb = DARK
    doc.styles["Subtitle"].font.size = Pt(12)
    doc.styles["Subtitle"].font.color.rgb = MUTED

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("旅游接待管理系统 | 甲方功能原型页面图册")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("第 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(footer, "PAGE")
    r = footer.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def add_field(paragraph, field: str):
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_title_page(doc: Document, item_count: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("旅游接待管理系统")
    set_run_font(r, size=28, color=DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("甲方功能原型页面图册")
    set_run_font(r, size=18, color=BLUE, bold=True)

    lines = [
        "结构：每个菜单页面先说明页面作用与功能，再贴对应的页面原型截图。",
        f"范围：共 {item_count} 个菜单页面，按 Excel 功能建设清单顺序生成。",
        "说明：截图为按业务需求绘制的原型页面，用于商务确认、需求沟通和范围评审。",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(line)
        set_run_font(r, size=12, color=DARK)


def add_page(doc: Document, item: dict[str, str], idx: int):
    if idx > 1:
        doc.add_page_break()
    title = item["优化建设内容"] or item["老系统菜单/功能"]

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{idx:02d}. {item['模块']} / {title}")
    set_run_font(r, size=17, color=BLUE, bold=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(5)
    r = meta.add_run(f"菜单路径：{item['模块']} / {item['老系统菜单/功能']}    {item['编号']}    {PHASE_TEXT.get(item['优先级'], item['优先级'])} / {item['建设方式']}")
    set_run_font(r, size=9.5, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("页面作用与功能")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(page_usage_summary(item))
    set_run_font(r, size=10, color=DARK)

    path = image_path_for(item, idx)
    if not path.exists():
        raise FileNotFoundError(path)
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_after = Pt(0)
    run = pic.add_run()
    run.add_picture(str(path), width=Inches(10.05))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(0)
    r = caption.add_run("页面原型截图")
    set_run_font(r, size=8.5, color=GREEN, bold=True)


def main():
    items = read_items()
    if len(items) != 67:
        raise RuntimeError(f"Expected 67 function items, got {len(items)}")
    if list(dict.fromkeys(item["模块"] for item in items)) != MODULE_ORDER:
        raise RuntimeError("Excel module order changed; please review generator.")

    images = sorted(IMAGE_DIR.glob("*.png"))
    if len(images) != 68:
        raise RuntimeError(f"Expected 68 PNG files, got {len(images)}")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_title_page(doc, len(items))
    for idx, item in enumerate(items, start=1):
        add_page(doc, item, idx)
    doc.core_properties.title = "旅游接待管理系统-甲方功能原型页面图册"
    doc.core_properties.subject = "菜单页面作用与原型截图"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
