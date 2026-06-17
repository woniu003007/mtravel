from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT_DIR = ROOT / "output" / "原型确认截图"
DOCX_PATH = ROOT / "文档" / "旅游接待管理系统-客户演示版功能说明.docx"

BASE_FONT = "Calibri"
CN_FONT = "Microsoft YaHei"
BLUE = RGBColor(37, 99, 235)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
TABLE_FILL = "EFF6FF"
SOFT_FILL = "F8FAFC"
GREEN_FILL = "ECFDF5"
WARN_FILL = "FFF7ED"
CONTENT_DXA = 9360


def clean(text: object) -> str:
    return str(text or "").replace("帐款", "账款").replace("帐号", "账号").strip()


def set_run_font(run, size=None, color=None, bold=None):
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def style_doc(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = BASE_FONT
        style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    title = doc.styles["Title"]
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = DARK
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = DARK
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("旅游接待管理系统 | 客户演示版")
    set_run_font(r, size=9, color=MUTED)


def add_para(doc, text="", *, style=None, size=None, color=None, bold=None, after=None):
    p = doc.add_paragraph(style=style)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(clean(text))
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths[min(i, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def cell_text(cell, text, *, bold=False, fill=None, color=None, size=9.4):
    if fill:
        shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(clean(text))
    set_run_font(r, size=size, color=color, bold=bold)


def add_callout(doc, title, body, fill=SOFT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(clean(title))
    set_run_font(r, size=10.5, color=BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(clean(body))
    set_run_font(r2, size=10, color=DARK)


def add_simple_table(doc, headers, rows, widths, *, header_fill=TABLE_FILL):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell_text(table.rows[0].cells[i], header, bold=True, fill=header_fill, color=BLUE)
    for r_i, row in enumerate(rows, start=1):
        for c_i, value in enumerate(row):
            cell_text(table.rows[r_i].cells[c_i], value, size=9.1)
    return table


def add_screenshot(doc, filename, caption):
    path = SCREENSHOT_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.55))
        add_para(doc, caption, size=8.8, color=MUTED, after=8)


def build():
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    add_para(doc, "给客户看的系统功能说明", color=BLUE, bold=True, size=10, after=2)
    add_para(doc, "旅游接待管理系统", style="Title")
    add_para(doc, "客户演示版功能说明", style="Subtitle")
    add_callout(
        doc,
        "一句话说明",
        "这个系统不是单独录订单的软件，而是把客户、订单、团队安排、导游报账、财务审核、应收应付和经营统计串成一条线，减少漏单、漏收款、重复排车、成本失控和对账困难。",
        fill=GREEN_FILL,
    )

    doc.add_heading("一、客户先看什么", level=1)
    add_simple_table(
        doc,
        ["先讲", "客户听得懂的说法", "客户要确认"],
        [
            ["为什么要做", "现在很多信息散在微信、Excel、人工记账里，系统要把关键业务统一管起来。", "现有痛点是否准确。"],
            ["每天怎么用", "销售录订单，计调排资源，导游报账，财务审核收付款，老板看进度和利润。", "岗位分工是否符合公司实际。"],
            ["一期先做什么", "先把会影响收款、排团、成本和审核的主流程做稳。", "一期范围是否认可。"],
            ["后续再做什么", "票务接口、AI客服、智能报价、知识库等放在后续增强。", "哪些可以后置。"],
        ],
        [1200, 5200, CONTENT_DXA - 6400],
    )

    doc.add_heading("二、用一条业务线讲系统", level=1)
    add_callout(
        doc,
        "推荐现场讲法",
        "不要从菜单逐个点。建议用一个真实团队举例：客户发团需求 → 销售确认订单 → 系统检查授信 → 计调安排导游/车/房/票 → 财务审核成本和应收 → 导游报账结算 → 老板看团队利润。",
    )
    add_simple_table(
        doc,
        ["步骤", "系统做什么", "客户能得到什么"],
        [
            ["1 客户建档", "保存客户资料、正式主体、合同、授信额度。", "下单、开票、收款对象不混乱。"],
            ["2 销售下单", "录订单、确认人数、确认金额、上传确认件。", "订单一确认，应收和团队人数同步生成。"],
            ["3 计调排团", "安排导游、车辆、酒店、景区、餐饮和外委地接。", "每个团还缺什么，一眼能看到。"],
            ["4 财务审核", "审核应收、应付、成本、利润、备用金和凭证。", "成本提前暴露，账款有来源。"],
            ["5 统计看板", "统计收客、利润、账龄、资源采购和团队进度。", "老板不用到处问，就能看经营情况。"],
        ],
        [900, 4300, CONTENT_DXA - 5200],
    )
    add_screenshot(doc, "01-业务工作台.png", "客户先看工作台：让客户知道系统每天提醒什么、谁要处理什么。")

    doc.add_heading("三、按岗位讲功能", level=1)
    add_simple_table(
        doc,
        ["岗位", "他们关心什么", "系统给他的页面"],
        [
            ["老板/管理层", "今天有多少订单没确认、多少团没排完、哪些客户快超授信、团队赚不赚钱。", "业务工作台、团队进度、利润统计、账款统计。"],
            ["销售", "客户能不能下单、订单确认件有没有、费用变更会不会影响应收。", "客户授信、订单管理、费用变更、电子合同、游客信息。"],
            ["计调", "导游、车、房、票、餐、地接安排完没有，有没有冲突和超预算。", "团队安排总控、计调审核、导游排班、房态库存、车调询价。"],
            ["财务", "应收应付从哪里来、收没收到、付没付出、凭证齐不齐、备用金有没有核销。", "财务审核、实时应收、应付管理、备用金闭环、导游结算、发票。"],
            ["系统管理员", "员工权限、审批规则、消息提醒、日志能不能管住。", "角色权限、员工管理、业务参数、审批预警、操作日志。"],
        ],
        [1200, 3900, CONTENT_DXA - 5100],
    )

    doc.add_heading("四、几个必须讲清楚的核心页面", level=1)
    core_pages = [
        ("02-客户授信.png", "客户授信：客户欠款和可下单额度要实时看见。", "客户发多个团时，避免到最后才发现欠款太多。"),
        ("03-订单确认.png", "订单管理：订单确认后自动带出应收、人数和团队关联。", "客户确认件、金额变更、取消订单都要有痕迹。"),
        ("04-计调安排.png", "团队安排总控：一个团的导、车、房、票、餐、地接集中看。", "计调不用在多个表里找状态。"),
        ("05-财务审核.png", "财务审核：团队收入、成本、毛利、异常凭证提前审核。", "财务不是事后算账，而是过程风控。"),
        ("06-实时应收.png", "实时应收：订单确认就形成应收，收款后释放授信。", "老板和财务随时知道谁欠款、欠多少、欠多久。"),
        ("07-收客统计.png", "收客统计：统一有效人数口径。", "避免绩效和经营统计口径不一致。"),
    ]
    for filename, title, body in core_pages:
        doc.add_heading(title, level=2)
        add_para(doc, body, color=DARK, bold=True)
        add_screenshot(doc, filename, f"页面截图：{title}")

    doc.add_heading("五、客户确认时只问这些问题", level=1)
    add_simple_table(
        doc,
        ["确认问题", "为什么要问"],
        [
            ["客户授信超了，是拦截还是走审批？", "决定销售能否继续下单，也影响财务风控。"],
            ["订单确认时，哪些附件必须上传？", "决定后续应收、合同和对账依据。"],
            ["团队安排里，哪些资源必须先排完才能提交审核？", "决定计调流程和异常提醒。"],
            ["备用金是给导游、计调，还是按团队统一申请？", "决定财务付款和核销流程。"],
            ["有效人数怎么统计？房差、儿童、老人、免票是否计入？", "决定统计和绩效口径。"],
            ["哪些功能一期必须上线，哪些可以二期？", "决定开发排期和验收边界。"],
        ],
        [3100, CONTENT_DXA - 3100],
        header_fill=WARN_FILL,
    )

    doc.add_heading("六、现场演示顺序", level=1)
    add_simple_table(
        doc,
        ["顺序", "打开页面", "讲给客户听的话"],
        [
            ["1", "业务工作台", "这是每天进系统先看的页面，告诉每个岗位今天要处理什么。"],
            ["2", "客户授信", "客户能不能继续下单，不靠人工记忆，系统自动看额度和欠款。"],
            ["3", "订单管理", "销售确认订单后，团队人数、应收、授信都会跟着变化。"],
            ["4", "团队安排总控", "计调看一个团的所有资源安排，不用多个 Excel 来回找。"],
            ["5", "财务团队审核", "财务提前审核成本、利润和凭证，问题及时退回。"],
            ["6", "实时应收", "每笔应收能追到订单和团队，收款后释放授信。"],
            ["7", "收客统计", "最后看经营结果：收了多少客、利润如何、账款风险在哪里。"],
        ],
        [700, 2100, CONTENT_DXA - 2800],
    )

    doc.add_heading("七、不要这样讲", level=1)
    add_callout(
        doc,
        "演示提醒",
        "不要一上来按菜单逐个念，也不要讲太多 P0/P1/P2、接口、字段、技术实现。客户最容易听懂的是：这个功能解决什么问题、谁来用、用了以后少做什么人工活、哪些规则需要他们确认。",
        fill=WARN_FILL,
    )

    props = doc.core_properties
    props.title = "旅游接待管理系统-客户演示版功能说明"
    props.subject = "客户功能演示"
    props.author = "Codex"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
